"""
AI-based extraction service for documents.
Uses Gemini 2.5 Flash for all extraction tasks (text and image extraction).
"""
import os
import re
import json
import base64
import io
import tempfile
from typing import Dict, Any, Optional
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from services.api.models.extraction import ExtractedField, FieldProvenance, DocumentExtractionResult
from services.api.services.text_quality_scorer import calculate_text_quality_score


def extract_with_ai(
    file_path: str,
    file_id: str,
    file_extension: str
) -> DocumentExtractionResult:
    """
    Extract data from document using AI-only approach.
    
    Process:
    1. Extract all text from document → Gemini 2.5 Flash for structured extraction
    2. Check text quality score - if low (< 0.5) or extraction fails, use vision extraction
    3. If document contains images → Extract images → Gemini 2.5 Flash for vision extraction
    4. Merge results from both text and image extraction
    
    Args:
        file_path: Path to document file
        file_id: Document UUID
        file_extension: File extension (e.g., '.pdf')
        
    Returns:
        DocumentExtractionResult with extracted fields and provenance
    """
    ext = file_extension.lower()
    
    # Check if file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document file not found: {file_path}")
    
    all_extracted_fields: Dict[str, ExtractedField] = {}
    raw_text = None
    page_texts = {}
    text_quality_score = None
    use_vision_fallback = False
    text_fields = {}  # Store text extraction results for potential fallback
    
    # Step 1: Extract text from document (if supported)
    if ext in ['.pdf', '.docx', '.doc', '.txt']:
        try:
            text_result = _extract_text_from_document(file_path, file_extension, file_id)
            if text_result:
                raw_text, page_texts, text_fields = text_result
                
                # Calculate text quality score
                if raw_text:
                    text_quality_score = calculate_text_quality_score(raw_text)
                    print(f"Text quality score: {text_quality_score}")
                    
                    # If quality is low (< 0.3), use vision extraction as fallback
                    if text_quality_score < 0.5:
                        print(f"Text quality score ({text_quality_score}) is low. Using vision extraction fallback.")
                        use_vision_fallback = True
                    else:
                        # Quality is acceptable, use text extraction results
                        all_extracted_fields.update(text_fields)
                else:
                    # No text extracted, use vision fallback
                    use_vision_fallback = True
            else:
                # Text extraction returned None, use vision fallback
                use_vision_fallback = True
        except Exception as e:
            print(f"Text extraction failed: {str(e)}. Using vision extraction fallback.")
            use_vision_fallback = True
    
    # Step 2: Use vision extraction if text quality is low or extraction failed
    if use_vision_fallback:
        try:
            page_images = _convert_document_to_images(file_path, file_extension)
            if page_images:
                print(f"Using vision extraction for {len(page_images)} page(s)")
                # Use Gemini 2.5 Flash for image extraction
                image_fields = _extract_with_vision_ai(page_images, file_id)
                # Vision extraction results override any text extraction results
                all_extracted_fields = image_fields
        except Exception as e:
            print(f"Vision extraction fallback failed: {str(e)}")
            # If vision also fails, we'll return whatever text results we have (if any)
            if not all_extracted_fields and text_fields:
                # Fall back to low-quality text results if vision fails
                all_extracted_fields.update(text_fields)
    
    # Step 3: Also extract images from PDF if it contains images (even if text quality is good)
    # This provides additional data from images embedded in PDFs
    if ext == '.pdf' and not use_vision_fallback:
        try:
            page_images = _extract_images_from_pdf(file_path)
            if page_images:
                # Use Gemini 2.5 Flash for image extraction
                image_fields = _extract_with_vision_ai(page_images, file_id)
                # Merge image fields (they may override text fields for better accuracy)
                all_extracted_fields.update(image_fields)
        except Exception as e:
            print(f"Image extraction from PDF failed: {str(e)}")
            # Continue with text-only results
    
    return DocumentExtractionResult(
        document_id=file_id,
        extraction_method="ai",
        extracted_fields=all_extracted_fields,
        raw_text=raw_text,
        page_texts=page_texts
    )


def _extract_text_from_document(
    file_path: str,
    file_extension: str,
    file_id: str
) -> Optional[tuple[str, dict[int, str], Dict[str, ExtractedField]]]:
    """
    Extract text from document and send to Gemini 2.5 Flash for structured extraction.
    
    Args:
        file_path: Path to document file
        file_extension: File extension
        file_id: Document UUID for provenance
    
    Returns:
        Tuple of (full_text, page_texts_dict, extracted_fields) or None if extraction fails
    """
    # Extract raw text based on file type
    full_text, page_texts = _get_document_text(file_path, file_extension)
    
    if not full_text or not full_text.strip():
        return None
    
        # Send to Gemini 2.5 Flash for structured extraction
    extracted_fields = _extract_with_gemini_text(full_text, file_id, page_texts)
    
    return (full_text, page_texts, extracted_fields)


def _get_document_text(file_path: str, file_extension: str) -> tuple[str, dict[int, str]]:
    """
    Extract raw text from document using appropriate library.
    
    Returns:
        Tuple of (full_text, page_texts_dict)
    """
    ext = file_extension.lower()
    
    if ext == '.pdf':
        return _extract_pdf_text(file_path)
    elif ext in ['.docx', '.doc']:
        return _extract_docx_text(file_path)
    elif ext == '.txt':
        return _extract_txt_text(file_path)
    else:
        return ("", {})


def _extract_pdf_text(file_path: str) -> tuple[str, dict[int, str]]:
    """Extract text from PDF using PyPDF2."""
    try:
        import PyPDF2
        
        full_text = []
        page_texts = {}
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        full_text.append(page_text)
                        page_texts[page_num] = page_text
                except Exception:
                    page_texts[page_num] = ""
        
        return "\n".join(full_text), page_texts
    
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF text extraction. Install with: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Failed to extract PDF text: {str(e)}")


def _extract_docx_text(file_path: str) -> tuple[str, dict[int, str]]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        full_text = []
        page_texts = {}
        
        # DOCX doesn't have clear page boundaries, so we'll treat it as one page
        page_text = "\n".join([para.text for para in doc.paragraphs])
        if page_text:
            full_text.append(page_text)
            page_texts[1] = page_text
        
        return "\n".join(full_text), page_texts
    
    except ImportError:
        raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Failed to extract DOCX text: {str(e)}")


def _extract_txt_text(file_path: str) -> tuple[str, dict[int, str]]:
    """Extract text from plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            return text, {1: text}
    except Exception as e:
        raise Exception(f"Failed to extract TXT text: {str(e)}")


def _extract_with_gemini_text(
    text: str,
    file_id: str,
    page_texts: dict[int, str]
) -> Dict[str, ExtractedField]:
    """
    Send extracted text to Gemini 2.5 Flash for structured extraction.
    
    Args:
        text: Full extracted text
        file_id: Document UUID (for provenance)
        page_texts: Dictionary of page_number -> text
        
    Returns:
        Dictionary of field_path -> ExtractedField
    """
    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("VISION_API_KEY")
    model_name = os.getenv("LLM_MODEL", "gemini-2.5-flash")
    
    if not api_key:
        return {}
    
    try:
        import google.genai as genai
        
        client = genai.Client(api_key=api_key)
        
        prompt = _get_text_extraction_prompt()
        
        # Limit text length for API (keep first 100k characters)
        text_for_api = text[:100000] if len(text) > 100000 else text
        
        response = client.models.generate_content(
            model=model_name,
            contents=f"{prompt}\n\nDocument Text:\n{text_for_api}"
        )
        
        response_text = response.text
        
        # Parse JSON from response
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            extracted_data = json.loads(json_match.group())
            return _convert_to_extracted_fields(extracted_data, file_id, page_texts, "text")
        else:
            return {}
    
    except ImportError:
        raise ImportError("google-generativeai library required. Install with: pip install google-generativeai")
    except (ConnectionError, OSError) as e:
        error_msg = str(e)
        if "getaddrinfo failed" in error_msg or "11001" in error_msg:
            raise Exception(f"Network connection failed: Cannot reach Gemini API. Check your internet connection and DNS settings. Error: {error_msg}")
        else:
            raise Exception(f"Network error during Gemini text extraction: {error_msg}")
    except Exception as e:
        error_msg = str(e)
        if "getaddrinfo failed" in error_msg or "11001" in error_msg:
            raise Exception(f"Network connection failed: Cannot reach Gemini API. Check your internet connection. Error: {error_msg}")
        raise Exception(f"Gemini text extraction failed: {error_msg}")


def _extract_images_from_pdf(file_path: str) -> Dict[int, bytes]:
    """
    Extract images from PDF pages (only if PDF contains images).
    Returns empty dict if no images found or if PDF is text-only.
    
    Returns:
        Dictionary mapping page_number -> image_bytes
    """
    try:
        from pdf2image import convert_from_path
        from pdf2image.exceptions import PDFInfoNotInstalledError
        import shutil
        
        # Get Poppler path from environment or try to find it
        poppler_path = os.getenv("POPPLER_PATH")
        
        # If not set, try common Windows locations
        if not poppler_path:
            common_paths = [
                r"C:\Program Files\poppler\Library\bin",
                r"C:\poppler\bin",
                r"C:\tools\poppler\bin",
                os.path.join(os.path.expanduser("~"), "poppler", "bin"),
            ]
            
            # Check if pdftoppm is in PATH
            if shutil.which("pdftoppm"):
                poppler_path = None  # Use system PATH
            else:
                # Try to find poppler in common locations
                for path in common_paths:
                    if os.path.exists(path) and os.path.exists(os.path.join(path, "pdftoppm.exe")):
                        poppler_path = path
                        break
        
        try:
            # Convert PDF pages to images
            if poppler_path:
                images = convert_from_path(file_path, dpi=200, poppler_path=poppler_path)
            else:
                images = convert_from_path(file_path, dpi=200)
            
            page_images = {}
            
            for page_num, image in enumerate(images, start=1):
                img_bytes = io.BytesIO()
                image.save(img_bytes, format='PNG')
                page_images[page_num] = img_bytes.getvalue()
            
            return page_images
        
        except PDFInfoNotInstalledError as e:
            error_msg = (
                f"Poppler is not found. Error: {str(e)}\n\n"
                "Solutions:\n"
                "1. Set POPPLER_PATH environment variable to Poppler bin directory\n"
                "2. Add Poppler bin directory to your system PATH\n"
                "3. Download from: https://github.com/oschwartz10612/poppler-windows/releases\n"
                "4. Or install via: conda install -c conda-forge poppler\n\n"
                f"Current PATH: {os.environ.get('PATH', 'Not set')[:200]}"
            )
            raise ValueError(error_msg)
    
    except ImportError:
        raise ImportError("pdf2image library required. Install with: pip install pdf2image")
    except Exception as e:
        raise Exception(f"Failed to extract images from PDF: {str(e)}")


def _convert_document_to_images(file_path: str, file_extension: str) -> Dict[int, bytes]:
    """
    Convert any document type to images for vision-based extraction.
    Used as fallback when text extraction fails or quality is low.
    
    Args:
        file_path: Path to document file
        file_extension: File extension (e.g., '.pdf', '.docx', '.txt')
        
    Returns:
        Dictionary mapping page_number -> image_bytes
    """
    ext = file_extension.lower()
    
    # For PDFs, use existing PDF to image conversion
    if ext == '.pdf':
        return _extract_images_from_pdf(file_path)
    
    # For DOCX files, try to convert to PDF first, then to images
    elif ext in ['.docx', '.doc']:
        try:
            # Try using docx2pdf if available
            try:
                from docx2pdf import convert
                
                # Create temporary PDF file
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_pdf:
                    tmp_pdf_path = tmp_pdf.name
                
                # Convert DOCX to PDF
                convert(file_path, tmp_pdf_path)
                
                # Convert PDF to images
                page_images = _extract_images_from_pdf(tmp_pdf_path)
                
                # Clean up temporary PDF
                try:
                    os.unlink(tmp_pdf_path)
                except:
                    pass
                
                return page_images
                
            except ImportError:
                # docx2pdf not available, try rendering text as image
                print("docx2pdf not available. Rendering DOCX text as image.")
                return _render_text_as_image(file_path, file_extension)
                
        except Exception as e:
            print(f"Failed to convert DOCX to images: {str(e)}. Rendering as text image.")
            return _render_text_as_image(file_path, file_extension)
    
    # For TXT files, render text as image
    elif ext == '.txt':
        return _render_text_as_image(file_path, file_extension)
    
    else:
        raise ValueError(f"Unsupported file format for vision extraction: {ext}")


def _render_text_as_image(file_path: str, file_extension: str) -> Dict[int, bytes]:
    """
    Render text file (TXT or DOCX) as an image for vision extraction.
    
    Args:
        file_path: Path to text file
        file_extension: File extension
        
    Returns:
        Dictionary mapping page_number -> image_bytes
    """
    try:
        # Read text content
        if file_extension.lower() in ['.docx', '.doc']:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        
        if not text or not text.strip():
            raise ValueError("No text content found in document")
        
        # Split text into pages (approximately 50 lines per page)
        lines = text.split('\n')
        lines_per_page = 50
        pages = []
        
        for i in range(0, len(lines), lines_per_page):
            page_text = '\n'.join(lines[i:i + lines_per_page])
            pages.append(page_text)
        
        if not pages:
            pages = [text]
        
        # Render each page as image
        page_images = {}
        page_width = 1200
        page_height = 1600
        margin = 50
        line_height = 30
        
        try:
            # Try to use a system font
            font = ImageFont.truetype("arial.ttf", 20)
        except:
            try:
                font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 20)
            except:
                # Fallback to default font
                font = ImageFont.load_default()
        
        for page_num, page_text in enumerate(pages, start=1):
            # Create image
            img = Image.new('RGB', (page_width, page_height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Draw text
            y = margin
            for line in page_text.split('\n'):
                if y + line_height > page_height - margin:
                    break
                draw.text((margin, y), line[:100], fill='black', font=font)
                y += line_height
            
            # Convert to bytes
            img_bytes = io.BytesIO()
            img.save(img_bytes, format='PNG')
            page_images[page_num] = img_bytes.getvalue()
        
        return page_images
        
    except Exception as e:
        raise Exception(f"Failed to render text as image: {str(e)}")


def _extract_with_vision_ai(page_images: Dict[int, bytes], file_id: str) -> Dict[str, ExtractedField]:
    """
    Use Gemini 2.5 Flash to extract structured data from document page images.
    
    Args:
        page_images: Dictionary of page_number -> image_bytes
        file_id: Document UUID for provenance
        
    Returns:
        Dictionary of field_path -> ExtractedField
    """
    extracted_fields = {}
    
    api_key = os.getenv("GEMINI_API_KEY") or os.getenv("VISION_API_KEY")
    model_name = os.getenv("IMAGE_VISION_MODEL", "gemini-2.5-flash")
    
    if not api_key:
        return extracted_fields
    
    try:
        import google.genai as genai
        
        client = genai.Client(api_key=api_key)
        
        prompt = _get_vision_extraction_prompt()
        
        # Process each page image
        for page_number, image_bytes in page_images.items():
            try:
                # Decode image and convert to base64 for API
                image = Image.open(io.BytesIO(image_bytes))
                
                # Convert PIL Image to base64
                img_buffer = io.BytesIO()
                image.save(img_buffer, format='PNG')
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
                
                # Call Gemini 2.5 Flash with image
                response = client.models.generate_content(
                    model=model_name,
                    contents=[
                        {"role": "user", "parts": [
                            {"text": prompt},
                            {"inline_data": {"mime_type": "image/png", "data": img_base64}}
                        ]}
                    ]
                )
                
                response_text = response.text
                
                # Parse JSON from response
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    extracted_data = json.loads(json_match.group())
                    page_fields = _convert_to_extracted_fields(
                        extracted_data, 
                        file_id, 
                        {page_number: ""}, 
                        "vision",
                        page_number
                    )
                    extracted_fields.update(page_fields)
            
            except Exception as e:
                print(f"Error processing page {page_number}: {str(e)}")
                continue
        
        return extracted_fields
    
    except ImportError:
        raise ImportError("google-genai library required. Install with: pip install google-genai")
    except Exception as e:
        raise Exception(f"Gemini vision extraction failed: {str(e)}")


def _convert_to_extracted_fields(
    extracted_data: Dict[str, Any],
    file_id: str,
    page_texts: dict[int, str],
    source_type: str,
    page_number: Optional[int] = None
) -> Dict[str, ExtractedField]:
    """
    Convert extracted data dictionary to ExtractedField format.
    
    Args:
        extracted_data: Dictionary from Gemini API response
        file_id: Document UUID
        page_texts: Page texts dictionary (for finding page numbers)
        source_type: "text" or "vision"
        page_number: Optional page number (for vision extraction)
        
    Returns:
        Dictionary of field_path -> ExtractedField
    """
    extracted_fields = {}
    
    # Flatten nested structure (e.g., {"location": {"street_address": {...}}})
    flattened = _flatten_extraction_data(extracted_data)
    
    for field_path, field_data in flattened.items():
        if isinstance(field_data, dict):
            value = field_data.get("value")
            confidence = field_data.get("confidence")
        else:
            value = field_data
            confidence = None
        
        if value is not None:
            # Determine page number
            if page_number:
                page_num = page_number
            elif page_texts:
                page_num = list(page_texts.keys())[0] if page_texts else None
            else:
                page_num = None
            
            extracted_fields[field_path] = ExtractedField(
                value=value,
                provenance=FieldProvenance(
                    file_id=file_id,
                    page_number=page_num,
                    source_type=source_type,
                    confidence=confidence
                )
            )
    
    return extracted_fields


def _flatten_extraction_data(data: Dict[str, Any], prefix: str = "") -> Dict[str, Any]:
    """
    Flatten nested extraction data structure.
    
    Example:
    {
      "location": {
        "street_address": {"value": "123 Main St", "confidence": 0.9}
      }
    }
    ->
    {
      "location.street_address": {"value": "123 Main St", "confidence": 0.9}
    }
    """
    flattened = {}
    
    for key, value in data.items():
        field_path = f"{prefix}.{key}" if prefix else key
        
        if isinstance(value, dict):
            if "value" in value or "confidence" in value:
                # This is a field with value/confidence
                flattened[field_path] = value
            else:
                # This is a nested section, recurse
                flattened.update(_flatten_extraction_data(value, field_path))
        else:
            # Simple value
            flattened[field_path] = value
    
    return flattened


def _get_text_extraction_prompt() -> str:
    """
    Get prompt for Gemini 2.5 Flash text extraction.
    """
    return """Extract structured MLS listing information from the following document text.

CRITICAL RULES:
- ONLY extract information that is CLEARLY PRESENT in the text
- DO NOT guess, infer, or fabricate values
- If a value is not present, return null
- Output MUST be valid JSON matching the provided schema

DATE FORMAT INSTRUCTIONS:
- All dates in the document will be in US format (MM/DD/YYYY), e.g., "04/02/2026" or "4/2/2026"
- Extract dates exactly as written in the document
- For expiration_date: You MUST output the date in date-time format (YYYY-MM-DDTHH:MM:SS)
- Convert US format dates to date-time format:
  * "04/02/2026" or "4/2/2026" → "2026-04-02T00:00:00"
  * "12/31/2025" → "2025-12-31T00:00:00"
  * Always use 4-digit year, 2-digit month, 2-digit day
  * Set time to 00:00:00 (midnight)
- If date is already in ISO format, keep it as date-time format

SPECIAL CONDITIONS EXTRACTION:
- special_conditions should ONLY be extracted if "short sale" is specifically mentioned in the document
- If "short sale" is mentioned, extract the relevant text about the short sale
- If "short sale" is NOT mentioned, return null (do not infer or guess)
- Do NOT extract other conditions unless they are explicitly about a short sale

PROPERTY CONDITION EXTRACTION:
- Check if the document indicates this is a "new construction" property
- Look for keywords like: "new construction", "new build", "newly constructed", "new home", "newly built", "under construction", "to be built", "pre-construction"
- If the document clearly indicates new construction, set property_condition to "new construction"
- If the document does NOT indicate new construction, set property_condition to "resale"
- Only extract if this information is clearly present in the document

WATERFRONT FEATURES EXTRACTION:
- waterfront_features should ONLY be extracted if the property is directly adjacent to a water body (lake, river, creek, pond, bay, ocean, etc.)
- If the document indicates the property is directly on or adjacent to water, extract the features of the water body:
  * Name of the water body (e.g., "Lake Travis", "Colorado River")
  * Type of water body (e.g., "Lake", "River", "Creek", "Pond", "Bay", "Ocean")
  * Any specific features mentioned (e.g., "sandy beach", "boat dock", "fishing access")
- Format: Combine name and type, e.g., "Lake Travis, Lake" or "Colorado River, River"
- If the property is NOT directly adjacent to water, return null
- Do NOT extract if only distance to water is mentioned (use distance_to_water field for that)

TAX INFORMATION EXTRACTION:
- When extracting tax information from tax documents, ALWAYS extract the LATEST/MOST RECENT year's tax data
- Look for multiple tax years in the document and identify the most recent one
- Extract tax_year, tax_annual_amount, tax_assessed_value, and tax_rate for the LATEST year only
- If multiple years are present (e.g., "2023: $5,000" and "2024: $5,500"), extract the higher/newer year (2024)
- If only one year is present, extract that year's information
- Do NOT extract older tax years - only the most recent/latest year

INTERMEDIARY EXTRACTION:
- First, check if intermediary information is already present in the document text (e.g., "Intermediary: Yes", "Intermediary: No", "Intermediary Status: Yes")
- If intermediary information is found in the text → extract that value (true/false)
- ONLY if intermediary information is NOT found in the text, then look for it in LISTING AGREEMENT documents:
  * Look for a section or field labeled "Intermediary" or "Intermediary Status" in the listing agreement
  * Check for a checkbox, box, or field that is:
    - Checked (✓, X, checkmark)
    - Crossed (X, ×)
    - Marked with any mark indicating selection
  * If a box/checkbox under "Intermediary" is checked, crossed, or marked → set intermediary to true
  * If the box is empty, unchecked, or not marked → set intermediary to false
  * If the "Intermediary" section/field is not present in the listing agreement → return null
- Priority: Text information first, then checkbox/box in listing agreement if not found

LIVING AREA EXTRACTION (CRITICAL - property.living_area_sqft):
- This is the INTERIOR HEATED/FINISHED square footage of the home, NOT lot size or garage size
- Extract the numeric value (integer) from ANY of these exact patterns you see in the document:
  * "Living Area: 2,500 sqft" → extract 2500
  * "Living Area Sqft: 2500" → extract 2500
  * "Total Living Area: 2,500" → extract 2500
  * "Heated Living Area: 2500 sq ft" → extract 2500
  * "Finished Living Area: 2,500 SF" → extract 2500
  * "Living Room Area: 2500" → extract 2500
  * "SFLA: 2500" or "SF LA: 2500" → extract 2500
  * "Square Feet: 2,500" (when context indicates living area) → extract 2500
  * "2,500 sqft" or "2500 sqft" (when labeled as living/heated/finished area) → extract 2500
- Look for these keywords near the number: "living", "heated", "finished", "interior", "SFLA", "SF LA"
- DO NOT extract: lot size, lot sqft, garage sqft, basement sqft (unless it's finished living space)
- Remove commas and extract ONLY the numeric value (e.g., "2,500" → 2500, "1,234.5" → 1234)
- If you see multiple living area values, prioritize in this order:
  1. "Heated Living Area" or "Heated Sqft"
  2. "Finished Living Area" or "Finished Sqft"
  3. "Total Living Area" or "Total Sqft"
  4. "Living Area" or "Living Sqft"
  5. Any other sqft value clearly indicating interior living space
- If the value includes decimals, round to nearest integer
- If units are in square meters, multiply by 10.764 to convert to square feet

Return a JSON object with the following structure. Each field should be an object with "value" and "confidence":

{
  "listing_meta": {
    "flex_listing": { "value": boolean | null, "confidence": number },
    "listing_agreement": { "value": string | null, "confidence": number },
    "listing_agreement_document": { "value": string | null, "confidence": number },
    "listing_service": { "value": string | null, "confidence": number },
    "list_price": { "value": number | null, "confidence": number },
    "expiration_date": { "value": string | null, "confidence": number },
    "special_conditions": { "value": string | null, "confidence": number }
  },
  "location": {
    "street_number": { "value": string | null, "confidence": number },
    "street_name": { "value": string | null, "confidence": number },
    "street_address": { "value": string | null, "confidence": number },
    "city": { "value": string | null, "confidence": number },
    "county": { "value": string | null, "confidence": number },
    "state": { "value": string | null, "confidence": number },
    "country": { "value": string | null, "confidence": number },
    "zip_code": { "value": string | null, "confidence": number },
    "subdivision": { "value": string | null, "confidence": number },
    "tax_legal_description": { "value": string | null, "confidence": number },
    "tax_lot": { "value": string | null, "confidence": number },
    "parcel_number": { "value": string | null, "confidence": number },
    "additional_parcel": { "value": boolean | null, "confidence": number },
    "additional_parcel_description": { "value": string | null, "confidence": number },
    "mla_area": { "value": string | null, "confidence": number },
    "flood_plain": { "value": boolean | null, "confidence": number },
    "etj": { "value": boolean | null, "confidence": number },
    "latitude": { "value": number | null, "confidence": number },
    "longitude": { "value": number | null, "confidence": number }
  },
  "schools": {
    "elementary_school_district": { "value": string | null, "confidence": number },
    "middle_junior_school": { "value": string | null, "confidence": number },
    "high_school": { "value": string | null, "confidence": number },
    "school_district": { "value": string | null, "confidence": number }
  },
  "property": {
    "property_sub_type": { "value": string | null, "confidence": number },
    "ownership_type": { "value": string | null, "confidence": number },
    "levels": { "value": number | null, "confidence": number },
    "main_level_bedrooms": { "value": number | null, "confidence": number },
    "other_level_bedrooms": { "value": number | null, "confidence": number },
    "year_built": { "value": number | null, "confidence": number },
    "year_built_source": { "value": string | null, "confidence": number },
    "bathrooms_full": { "value": number | null, "confidence": number },
    "bathrooms_half": { "value": number | null, "confidence": number },
    "living_area_sqft": { "value": number | null, "confidence": number },
    "living_area_source": { "value": string | null, "confidence": number },
    "garage_spaces": { "value": number | null, "confidence": number },
    "parking_total": { "value": number | null, "confidence": number },
    "direction_faces": { "value": string | null, "confidence": number },
    "lot_size_acres": { "value": number | null, "confidence": number },
    "property_condition": { "value": string | null, "confidence": number },
    "view": { "value": string | null, "confidence": number },
    "distance_to_water": { "value": number | null, "confidence": number },
    "waterfront_features": { "value": string | null, "confidence": number },
    "restrictions": { "value": string | null, "confidence": number },
    "living_room": { "value": string | null, "confidence": number },
    "dining_room": { "value": string | null, "confidence": number },
    "construction_material": { "value": string[] | [], "confidence": number },
    "foundation_details": { "value": string[] | [], "confidence": number },
    "roof": { "value": string[] | [], "confidence": number },
    "lot_features": { "value": string[] | [], "confidence": number }
  },
  "features": {
    "interior_features": { "value": string[] | [], "confidence": number },
    "exterior_features": { "value": string[] | [], "confidence": number },
    "patio_porch_features": { "value": string[] | [], "confidence": number },
    "fireplaces": { "value": string[] | [], "confidence": number },
    "flooring": { "value": string[] | [], "confidence": number },
    "accessibility_features": { "value": string[] | [], "confidence": number },
    "horse_amenities": { "value": string[] | [], "confidence": number },
    "other_structures": { "value": string[] | [], "confidence": number },
    "appliances": { "value": string[] | [], "confidence": number },
    "pool_features": { "value": string[] | [], "confidence": number },
    "guest_accommodations": { "value": string | null, "confidence": number },
    "window_features": { "value": string[] | [], "confidence": number },
    "security_features": { "value": string[] | [], "confidence": number },
    "laundry_location": { "value": string | null, "confidence": number },
    "fencing": { "value": string | null, "confidence": number },
    "community_features": { "value": string[] | [], "confidence": number }
  },
  "utilities": {
    "utilities": { "value": string[] | [], "confidence": number },
    "heating": { "value": string[] | [], "confidence": number },
    "cooling": { "value": string[] | [], "confidence": number },
    "water_source": { "value": string[] | [], "confidence": number },
    "sewer": { "value": string[] | [], "confidence": number },
    "documents_available": { "value": string[] | [], "confidence": number },
    "disclosures": { "value": string[] | [], "confidence": number }
  },
  "green_energy": {
    "green_energy": { "value": string[] | [], "confidence": number },
    "green_sustainability": { "value": string[] | [], "confidence": number }
  },
  "financial": {
    "association": { "value": boolean | null, "confidence": number },
    "association_name": { "value": string | null, "confidence": number },
    "association_fee": { "value": number | null, "confidence": number },
    "association_amount": { "value": number | null, "confidence": number },
    "acceptable_financing": { "value": string[] | [], "confidence": number },
    "estimated_tax": { "value": number | null, "confidence": number },
    "tax_year": { "value": number | null, "confidence": number },
      // NOTE: Extract the LATEST/MOST RECENT year from tax documents. If multiple years present, use the newest year.
    "tax_annual_amount": { "value": number | null, "confidence": number },
      // NOTE: Extract for the latest tax year only.
    "tax_assessed_value": { "value": number | null, "confidence": number },
      // NOTE: Extract for the latest tax year only.
    "tax_rate": { "value": number | null, "confidence": number },
    "buyer_incentive": { "value": string | null, "confidence": number },
    "tax_exemptions": { "value": string[] | [], "confidence": number },
    "possession": { "value": string | null, "confidence": number },
    "seller_contributions": { "value": boolean | null, "confidence": number },
    "intermediary": { "value": boolean | null, "confidence": number }
      // NOTE: First check text for intermediary info. If not found, then check listing agreement for checked/crossed box under "Intermediary" title. If checked/marked → true, if empty/unchecked → false, if not present → null.
  },
  "showing": {
    "occupant_type": { "value": string | null, "confidence": number },
    "showing_requirements": { "value": string[] | [], "confidence": number },
    "owner_name": { "value": string | null, "confidence": number },
    "lockbox_type": { "value": string | null, "confidence": number },
    "lockbox_location": { "value": string | null, "confidence": number },
    "showing_instructions": { "value": string | null, "confidence": number }
  },
  "agents": {
    "listing_agent": { "value": string | null, "confidence": number },
    "co_listing_agent": { "value": string | null, "confidence": number }
  },
  "remarks": {
    "directions": { "value": string | null, "confidence": number },
    "private_remarks": { "value": string | null, "confidence": number },
    "public_remarks": { "value": string | null, "confidence": number },
    "syndication_remarks": { "value": string | null, "confidence": number }
  }
}

Confidence guidelines:
- 0.9–1.0 → clearly present, unambiguous
- 0.6–0.8 → present but minor ambiguity
- <0.6 → weak presence (still do NOT guess)"""


def _get_vision_extraction_prompt() -> str:
    """
    Get prompt for Gemini 2.5 Flash vision extraction (same schema as text extraction).
    """
    # Use the same detailed prompt as text extraction but with vision-specific instructions
    base_prompt = _get_text_extraction_prompt()
    
    vision_instructions = """You are analyzing an IMAGE of a real estate document.
Extract structured MLS listing information from what you can SEE in the image.

CRITICAL RULES FOR VISION:
- ONLY extract information that is CLEARLY VISIBLE in the image
- DO NOT guess, infer, or fabricate values
- If a value is not visible or not legible, return null
- DO NOT perform OCR-style full text reconstruction
- Focus on structured fields that are clearly visible
- This data will be reviewed and edited by a human before MLS submission

"""
    
    return vision_instructions + base_prompt
